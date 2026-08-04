from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell

from app.models import MonitoringItem, Project


def export_project_docx(
    project: Project,
    items: list[MonitoringItem],
    template_path: Path,
    output_path: Path,
    export_type: str,
) -> Path:
    doc = Document(str(template_path))
    if export_type == "basic":
        _fill_basic(doc, project, items)
    else:
        _fill_annual(doc, project, items)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _set_cell_text(cell: _Cell, text: str) -> None:
    text = text or ""
    # clear paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
            return
        p.add_run(text)
        return


def _clear_table_keep_header(table: Table, header_rows: int = 1) -> None:
    while len(table.rows) > header_rows:
        row = table.rows[-1]
        row._tr.getparent().remove(row._tr)


def _add_row(table: Table) -> list[_Cell]:
    new_tr = deepcopy(table.rows[-1]._tr)
    # clear text in cloned row
    for tc in new_tr.findall(qn("w:tc")):
        for t in tc.iter(qn("w:t")):
            t.text = ""
    table._tbl.append(new_tr)
    return table.rows[-1].cells


def _fill_paragraph_title(doc: Document, project: Project) -> None:
    title = project.name or project.client_name or "监测方案"
    if project.project_type == "annual" and "方案" not in title:
        title = f"{title}自行监测方案"
    for p in doc.paragraphs[:5]:
        if p.text.strip():
            # replace whole first non-empty paragraph as title-ish
            if "方案" in p.text or "公司" in p.text or "监测" in p.text or p.text.startswith("xxx"):
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = title
                else:
                    p.add_run(title)
                break


def _fill_annual(doc: Document, project: Project, items: list[MonitoringItem]) -> None:
    _fill_paragraph_title(doc, project)
    if not doc.tables:
        raise ValueError("年度模板中未找到表格")
    table = doc.tables[0]
    _clear_table_keep_header(table, 1)

    # ensure at least one data row template exists after header
    # after clear, only header remains — clone header structure for new rows via add
    # python-docx: add_row copies last row style
    # We need one blank data row first
    if len(table.rows) == 1:
        table.add_row()

    # remove the auto blank if we will rewrite — keep one skeleton
    skeleton = table.rows[1]
    # write rows
    first = True
    last_category = None
    for it in items:
        if first:
            row_cells = skeleton.cells
            first = False
        else:
            row_cells = _add_row(table)

        category = it.category or ""
        # display blank category when same as previous for readability
        display_cat = category if category != last_category else ""
        last_category = category or last_category

        values = [
            display_cat,
            it.outlet_code or "",
            it.outlet_name or it.point_location or "",
            it.factors or "",
            it.sample_freq or it.period_freq or "",
            it.annual_times or it.period_freq or "",
            it.remark or "",
        ]
        for idx, val in enumerate(values):
            if idx < len(row_cells):
                _set_cell_text(row_cells[idx], val)

    if first:
        # no items — leave one empty data row
        pass


def _fill_basic(doc: Document, project: Project, items: list[MonitoringItem]) -> None:
    # Title
    for p in doc.paragraphs[:3]:
        if "方案" in p.text or p.text.startswith("xxx"):
            title = project.name or f"{project.client_name}环境检测方案"
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = title
            else:
                p.add_run(title)
            break

    if not doc.tables:
        raise ValueError("基础模板中未找到表格")

    # Table 0: project basics
    basic = doc.tables[0]
    mapping = {
        "项目名称": project.name,
        "建设单位": project.client_name,
        "建设地点": project.address,
        "经纬度": " / ".join(x for x in [project.longitude, project.latitude] if x) or "/",
        "项目概况": project.overview or "/",
        "建设单位联系人": project.contact,
        "建设单位联系方式": project.phone,
    }
    for row in basic.rows:
        cells = row.cells
        texts = [c.text.strip() for c in cells]
        for i, t in enumerate(texts):
            if t in mapping and i + 1 < len(cells):
                # avoid writing into merged duplicate
                target = cells[i + 1]
                # if label is 联系人 style with 5 cells
                _set_cell_text(target, mapping[t] or "")
            # special 5-col row: 联系人 / 联系方式
            if t == "建设单位联系人" and len(cells) >= 5:
                _set_cell_text(cells[1], project.contact or "")
                _set_cell_text(cells[3], project.phone or "")

    # categorize items
    unorg = [i for i in items if _is_cat(i, ("无组织",))]
    waste = [i for i in items if _is_cat(i, ("废水", "污水", "雨水"))]
    noise = [i for i in items if _is_cat(i, ("噪声", "噪音"))]
    org = [i for i in items if _is_cat(i, ("有组织", "废气")) and not _is_cat(i, ("无组织",))]
    # if unorg empty and org took 废气 only — already handled

    # tables: 1 unorg, 2 waste, 3 noise, 4 org (based on template analysis)
    if len(doc.tables) > 1:
        _fill_factor_table(
            doc.tables[1],
            unorg,
            cols=("category", "outlet_code", "point", "factors", "days", "freq", "remark"),
            defaults_category="无组织",
        )
    if len(doc.tables) > 2:
        _fill_factor_table(
            doc.tables[2],
            waste,
            cols=("outlet_code", "point", "factors", "days", "freq"),
        )
    if len(doc.tables) > 3:
        _fill_factor_table(
            doc.tables[3],
            noise,
            cols=("outlet_code", "point", "days", "freq"),
        )
    if len(doc.tables) > 4:
        _fill_factor_table(
            doc.tables[4],
            org,
            cols=("category", "outlet_code", "point", "factors", "days", "freq", "remark"),
            defaults_category="有组织",
        )


def _is_cat(item: MonitoringItem, keys: tuple[str, ...]) -> bool:
    blob = f"{item.category}{item.outlet_name}{item.point_location}"
    return any(k in blob for k in keys)


def _fill_factor_table(
    table: Table,
    items: list[MonitoringItem],
    cols: tuple[str, ...],
    defaults_category: str = "",
) -> None:
    # keep header row 0; drop other rows except trailing standard row if present
    header = [c.text.strip() for c in table.rows[0].cells]
    # identify standard rows (执行标准)
    standard_rows_xml = []
    data_start = 1
    for idx in range(1, len(table.rows)):
        row_text = "".join(c.text for c in table.rows[idx].cells)
        if "执行标准" in row_text:
            standard_rows_xml.append(deepcopy(table.rows[idx]._tr))

    _clear_table_keep_header(table, 1)

    if not items:
        # restore one empty + standards
        table.add_row()
        for tr in standard_rows_xml:
            table._tbl.append(tr)
        return

    # need a template row: add_row from header may be wrong width — use add_row anyway
    first = True
    for it in items:
        if first:
            table.add_row()
            cells = table.rows[-1].cells
            first = False
        else:
            cells = _add_row(table)

        point = it.point_location or it.outlet_name or ""
        days = it.monitor_days or ""
        freq = it.sample_freq or it.period_freq or ""
        mapping = {
            "category": it.category or defaults_category,
            "outlet_code": it.outlet_code or "",
            "point": point,
            "factors": it.factors or "",
            "days": days,
            "freq": freq,
            "remark": it.remark or "",
        }
        for i, key in enumerate(cols):
            if i < len(cells):
                _set_cell_text(cells[i], mapping.get(key, ""))

    for tr in standard_rows_xml:
        table._tbl.append(tr)


def safe_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', "_", name).strip() or "export"
    return name[:80]
